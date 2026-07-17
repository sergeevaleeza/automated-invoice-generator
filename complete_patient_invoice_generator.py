#!/usr/bin/env python3
"""
Patient Invoice & Cover Letter Generator
Generates per-patient invoices (PDF) and cover letters (DOCX) from patient roster and invoice data.
"""

import pandas as pd
import numpy as np
from pathlib import Path
import json
import re
from datetime import datetime, timedelta
from typing import Dict, List, Tuple, Optional, Set
import logging
import os
from difflib import SequenceMatcher
from io import BytesIO

# PDF generation
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_RIGHT, TA_CENTER
from reportlab.lib.utils import ImageReader

# DOCX handling
from docx import Document

# Shared data structures + formatting helpers (also used by excel_invoice_generator.py)
from invoice_models import (
    PatientData, InvoiceLine, ProcessingSummary, ProcessedPatientRecord, format_date_for_display,
    REQUIRED_TEMPLATE_PLACEHOLDERS, ValidationIssue, ValidationReport, VALIDATION_CATEGORIES,
    extract_embedded_cpt_code, SuperbillServiceLine,
    NOTICE_LEVEL_NORMAL, NOTICE_LEVEL_SECOND, NOTICE_LEVEL_FINAL, NOTICE_LEVEL_TITLES,
)
from clinic_config import load_clinic_config
import run_history
from qr_code import resolve_qr_image_bytes

# Excel invoice generation (mirrors the PDF layout, no shared business logic)
from excel_invoice_generator import generate_excel_invoice

class PatientInvoiceGenerator:
    """Main class for generating patient invoices and cover letters"""

    # Layout parameter tiers applied in order for single-page fitting
    _LAYOUT_TIERS = [
        # Tier 0: Base settings
        dict(spacer_header=8, spacer_sections=20, spacer_small=20,
             table_top_pad=6, table_bot_pad=6,
             font_body=9, font_header=10, font_header2=9, font_title=11, font_table=8),
        # Tier 1: Reduce spacers by ~25%
        dict(spacer_header=6, spacer_sections=15, spacer_small=15,
             table_top_pad=6, table_bot_pad=6,
             font_body=9, font_header=10, font_header2=9, font_title=11, font_table=8),
        # Tier 2: Reduce table padding from 6 to 3
        dict(spacer_header=6, spacer_sections=15, spacer_small=15,
             table_top_pad=3, table_bot_pad=3,
             font_body=9, font_header=10, font_header2=9, font_title=11, font_table=8),
        # Tier 3: Reduce body font from 9 to 8
        dict(spacer_header=6, spacer_sections=15, spacer_small=15,
             table_top_pad=3, table_bot_pad=3,
             font_body=8, font_header=10, font_header2=9, font_title=11, font_table=7),
        # Tier 4: Reduce header font from 10 to 9
        dict(spacer_header=6, spacer_sections=15, spacer_small=15,
             table_top_pad=3, table_bot_pad=3,
             font_body=8, font_header=9, font_header2=8, font_title=10, font_table=7),
        # Tier 5: Last resort — reduce everything further
        dict(spacer_header=4, spacer_sections=10, spacer_small=8,
             table_top_pad=2, table_bot_pad=2,
             font_body=7, font_header=9, font_header2=8, font_title=10, font_table=6),
    ]

    # A fuzzy match below this score is flagged for review during pre-flight
    # validation, even though it's already above the harder 0.6-per-name-part
    # cutoff _match_patient uses to consider it a candidate at all. Matches
    # this constants value against the existing ambiguity threshold in
    # _match_patient (also 0.85) for a consistent "high confidence" bar.
    LOW_CONFIDENCE_THRESHOLD = 0.85

    # Days after a 2nd Notice before suggest_notice_level() escalates to a
    # Final Notice — matches the "contact us within 14 days" language in
    # both notice letter templates.
    NOTICE_ESCALATION_DAYS = 14

    # Reminder letter templates for escalated notices, keyed by NOTICE_LEVEL_*
    # (replaces the normal cover letter for these patients). Converted from
    # the clinic's real Word mail-merge .doc originals (kept alongside, for
    # reference) to single-record .docx files using this app's [Bracket]
    # placeholder convention instead of Word MERGEFIELD codes.
    NOTICE_TEMPLATE_FILES = {
        NOTICE_LEVEL_SECOND: Path(__file__).parent / "templates" / "TEMPLATE_MAIL_MERGE_1st_level_YYYYMMDD.docx",
        NOTICE_LEVEL_FINAL: Path(__file__).parent / "templates" / "TEMPLATE_MAIL_MERGE_2nd_level_YYYYMMDD.docx",
    }

    def __init__(self, amount_due_strategy: str = "auto", statement_date: Optional[str] = None,
                 clinic_config: Optional[Dict] = None):
        self.amount_due_strategy = amount_due_strategy
        self.statement_date = datetime.strptime(statement_date, "%Y-%m-%d") if statement_date else datetime.now()
        self.payment_due_date = self._calculate_payment_due_date()
        # Raises ClinicConfigError (caught upstream, shown via st.error) if
        # clinic_config.json is missing/incomplete — fail fast rather than
        # generate invoices with wrong or placeholder clinic identity.
        self.clinic = clinic_config if clinic_config is not None else load_clinic_config()

        # Column mapping for normalization
        self.column_aliases = {
            'name': ["Name", "Patient Name", "[LastName, FirstName]"],
            'visit_date': ["Visit Date", "Service Date", "DOS", "Date of Service"],
            'total_amount': ["Total amount", "Charge", "Billed Amount", "Total Charge"],
            'copay': ["Copay", "Co-pay", "Copayment"],
            'paid': ["Paid", "Patient Paid", "Payments"],
            'previous_balance': ["Previous Balance", "Outstanding Balance", "Prior Balance", "Carryover"],
            'insurance': ["Insurance"],
            'type_of_service': ["Type Of Service", "Service Type", "Description", "Service Description"],
            'cpt_code': ["CPT Code", "CPT", "Procedure Code"],
            'icd10_code': ["ICD-10", "ICD10", "Diagnosis Code", "ICD Code"],
        }

        # Setup logging
        logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
        self.logger = logging.getLogger(__name__)

    def _calculate_payment_due_date(self) -> datetime:
        """Calculate payment due date (1 month from statement date, adjusted for weekends)"""
        due_date = self.statement_date + timedelta(days=30)

        # If weekend, move to next Monday
        if due_date.weekday() == 5:  # Saturday
            due_date += timedelta(days=2)
        elif due_date.weekday() == 6:  # Sunday
            due_date += timedelta(days=1)

        return due_date

    def _normalize_column_name(self, df: pd.DataFrame, target_column: str, custom_mapping: Dict = None) -> Optional[str]:
        """Find the actual column name using aliases"""
        if custom_mapping and target_column in custom_mapping:
            if custom_mapping[target_column] in df.columns:
                return custom_mapping[target_column]

        if target_column in self.column_aliases:
            for alias in self.column_aliases[target_column]:
                if alias in df.columns:
                    return alias

        # Direct match
        if target_column in df.columns:
            return target_column

        return None

    def _calculate_string_similarity(self, str1: str, str2: str) -> float:
        """Calculate similarity between two strings using sequence matching"""
        if not str1 or not str2:
            return 0.0

        # Normalize strings
        s1 = str1.lower().strip()
        s2 = str2.lower().strip()

        if s1 == s2:
            return 1.0

        # Use SequenceMatcher for similarity
        similarity = SequenceMatcher(None, s1, s2).ratio()

        # Boost score for very close matches (like "Shley" vs "Shely")
        if similarity >= 0.8:
            return similarity

        # Check for substring matches
        if s1 in s2 or s2 in s1:
            return max(similarity, 0.7)

        # Check for common prefixes/suffixes
        if len(s1) >= 3 and len(s2) >= 3:
            if s1[:3] == s2[:3] or s1[-3:] == s2[-3:]:
                return max(similarity, 0.6)

        return similarity

    def _sanitize_filename(self, name: str) -> str:
        """Sanitize filename to only contain A-Z, a-z, 0-9, _"""
        return re.sub(r'[^A-Za-z0-9_]', '_', name)

    def _has_cpt_codes(self, patient_df: pd.DataFrame) -> bool:
        """Check if any value in type_of_service column contains a 5-digit CPT code.
        Matches both bare codes ('90837') and codes embedded in descriptions
        ('Med Management (CPT Code 99213)').
        """
        if 'type_of_service' not in patient_df.columns:
            return False
        return any(extract_embedded_cpt_code(v) for v in patient_df['type_of_service'])

    def _count_pdf_pages(self, pdf_bytes: bytes) -> int:
        """Count pages in ReportLab-generated PDF by counting /Page objects"""
        return len(re.findall(rb'/Type\s*/Page\b', pdf_bytes))

    def _parse_patient_name(self, name: str) -> Tuple[str, str]:
        """Parse 'LastName, FirstName' format with support for complex names"""
        if ',' in name:
            parts = name.split(',', 1)
            last_name = parts[0].strip()
            first_name = parts[1].strip()

            # Handle complex last names like "Russell (Kwon)"
            if '(' in last_name:
                main_last_name = last_name.split('(')[0].strip()
                return first_name, main_last_name

            return first_name, last_name
        else:
            # Fallback: assume single name is last name
            return "", name.strip()

    def load_patient_roster(self, roster_file: str) -> Dict[str, PatientData]:
        """Load patient roster CSV and create lookup dictionary"""
        try:
            df = pd.read_csv(roster_file)
            self.logger.info(f"Loaded roster with {len(df)} patients")
            self.logger.info(f"Roster columns: {list(df.columns)}")

            patients = {}

            # Try to identify columns dynamically
            possible_prn_cols = ['Patient Identifier', 'Patient Record Number', 'PRN', 'ID', 'Patient ID']
            possible_first_cols = ['Patient First Name', 'First name', 'First Name', 'FirstName', 'Given Name']
            possible_last_cols = ['Patient Last Name', 'Last name', 'Last Name', 'LastName', 'Surname', 'Family Name']
            possible_dob_cols = ['DOB', 'Date of Birth', 'Birth Date']
            possible_addr1_cols = ['Address Line 1', 'Address 1', 'Street Address', 'Address']
            possible_addr2_cols = ['Address Line 2', 'Address 2', 'Apt', 'Suite']
            possible_city_cols = ['City']
            possible_state_cols = ['State', 'ST']
            possible_zip_cols = ['Postal Code', 'Zip Code', 'ZIP', 'Zip']

            def find_column(df, possible_names):
                for name in possible_names:
                    if name in df.columns:
                        return name
                return None

            # Standard CSV processing
            prn_col = find_column(df, possible_prn_cols)
            first_col = find_column(df, possible_first_cols)
            last_col = find_column(df, possible_last_cols)
            dob_col = find_column(df, possible_dob_cols)
            addr1_col = find_column(df, possible_addr1_cols)
            addr2_col = find_column(df, possible_addr2_cols)
            city_col = find_column(df, possible_city_cols)
            state_col = find_column(df, possible_state_cols)
            zip_col = find_column(df, possible_zip_cols)

            # Log what columns were found
            self.logger.info(f"Detected columns - PRN: {prn_col}, First: {first_col}, Last: {last_col}")

            for _, row in df.iterrows():
                # Clean up postal code to remove decimal formatting
                postal_code = row.get(zip_col, '') if zip_col else ''
                if pd.notnull(postal_code):
                    # Convert to string and remove .0 if present
                    postal_code = str(postal_code)
                    if postal_code.endswith('.0'):
                        postal_code = postal_code[:-2]
                    # Also remove any other decimal patterns
                    if '.' in postal_code and postal_code.replace('.', '').isdigit():
                        postal_code = postal_code.split('.')[0]
                else:
                    postal_code = ''

                patient = PatientData(
                    prn=str(row.get(prn_col, '')) if prn_col and pd.notnull(row.get(prn_col)) else '',
                    first_name=str(row.get(first_col, '')) if first_col and pd.notnull(row.get(first_col)) else '',
                    last_name=str(row.get(last_col, '')) if last_col and pd.notnull(row.get(last_col)) else '',
                    dob=str(row.get(dob_col, '')) if dob_col and pd.notnull(row.get(dob_col)) else '',
                    address_line1=str(row.get(addr1_col, '')) if addr1_col and pd.notnull(row.get(addr1_col)) else '',
                    address_line2=str(row.get(addr2_col, '')) if addr2_col and pd.notnull(row.get(addr2_col)) else '',
                    city=str(row.get(city_col, '')) if city_col and pd.notnull(row.get(city_col)) else '',
                    state=str(row.get(state_col, '')) if state_col and pd.notnull(row.get(state_col)) else '',
                    postal_code=postal_code
                )

                # Create lookup keys
                if patient.prn and patient.prn != 'nan':
                    patients[f"prn_{patient.prn}"] = patient

                name_key = f"{patient.first_name.lower()}_{patient.last_name.lower()}"
                if name_key not in patients:  # First match wins
                    patients[name_key] = patient

            self.logger.info(f"Loaded {len(patients)} unique patients into lookup")
            return patients

        except Exception as e:
            self.logger.error(f"Error loading patient roster: {e}")
            raise

    def _format_date_for_display(self, date_value) -> str:
        """Format date value to MM/DD/YYYY format"""
        return format_date_for_display(date_value, self.logger)

    def load_invoice_data(self, invoice_file: str, custom_mapping: Dict = None) -> pd.DataFrame:
        """Load invoice Excel file and normalize columns"""
        try:
            df = pd.read_excel(invoice_file, sheet_name='Sheet1')
            self.logger.info(f"Loaded invoice data from Sheet1 with {len(df)} rows")

            # Normalize column names
            column_mapping = {}
            required_columns = ['name', 'visit_date', 'total_amount', 'copay', 'paid']

            for col in required_columns:
                actual_col = self._normalize_column_name(df, col, custom_mapping)
                if actual_col is None:
                    raise ValueError(f"Required column '{col}' not found. Available columns: {list(df.columns)}")
                column_mapping[actual_col] = col

            # Optional columns
            optional_columns = ['previous_balance', 'insurance', 'type_of_service', 'cpt_code', 'icd10_code']
            for col in optional_columns:
                actual_col = self._normalize_column_name(df, col, custom_mapping)
                if actual_col:
                    column_mapping[actual_col] = col

            # Rename columns
            df = df.rename(columns=column_mapping)

            # Fill missing optional columns
            if 'previous_balance' not in df.columns:
                df['previous_balance'] = 0
            if 'insurance' not in df.columns:
                df['insurance'] = ''
            if 'type_of_service' not in df.columns:
                df['type_of_service'] = ''
            if 'cpt_code' not in df.columns:
                df['cpt_code'] = ''
            if 'icd10_code' not in df.columns:
                df['icd10_code'] = ''

            # Clean and convert data types
            df['total_amount'] = pd.to_numeric(df['total_amount'], errors='coerce').fillna(0)
            df['copay'] = pd.to_numeric(df['copay'], errors='coerce').fillna(0)
            df['paid'] = pd.to_numeric(df['paid'], errors='coerce').fillna(0)
            df['previous_balance'] = pd.to_numeric(df['previous_balance'], errors='coerce').fillna(0)

            # Clean text columns
            df['type_of_service'] = df['type_of_service'].fillna('').astype(str)
            df['cpt_code'] = df['cpt_code'].fillna('').astype(str).str.strip()
            df['icd10_code'] = df['icd10_code'].fillna('').astype(str).str.strip()

            return df

        except Exception as e:
            self.logger.error(f"Error loading invoice data: {e}")
            raise

    def _calculate_amount_due(self, row: pd.Series) -> float:
        """Calculate amount due based on strategy"""
        total_amount = float(row['total_amount'])
        copay = float(row['copay'])
        paid = float(row['paid'])

        if self.amount_due_strategy == "auto":
            if copay > 0:
                amount_due = copay - paid
            else:
                amount_due = total_amount - paid
        elif self.amount_due_strategy == "copay_minus_paid":
            amount_due = copay - paid
        elif self.amount_due_strategy == "total_minus_paid":
            amount_due = total_amount - paid
        else:
            raise ValueError(f"Unknown amount due strategy: {self.amount_due_strategy}")

        return max(0, amount_due)  # Floor at 0

    def _match_patient(self, name: str, patients: Dict[str, PatientData]) -> Tuple[Optional[PatientData], bool, float]:
        """Match patient by name with fuzzy matching support.

        Returns (patient_or_None, is_ambiguous, confidence_score) where
        confidence_score is 1.0 for an exact match, the best fuzzy-match
        score in [0, 1) for a fuzzy match, or 0.0 if no match was found.
        """
        first_name, last_name = self._parse_patient_name(name)

        # Try exact match first
        name_key = f"{first_name.lower()}_{last_name.lower()}"
        if name_key in patients:
            return patients[name_key], False, 1.0

        # Try partial/fuzzy matching with string similarity
        first_lower = first_name.lower().strip()
        last_lower = last_name.lower().strip()

        # Extract name components for better matching
        first_parts = [part.strip() for part in first_lower.replace(',', ' ').split() if part.strip()]
        last_parts = [part.strip() for part in last_lower.replace(',', ' ').split() if part.strip()]

        matches = []

        for key, patient in patients.items():
            if key.startswith('prn_'):
                continue

            patient_first = patient.first_name.lower().strip()
            patient_last = patient.last_name.lower().strip()

            # Calculate similarity scores
            first_name_score = 0
            last_name_score = 0

            # Score first name matching
            if first_parts and patient_first:
                max_first_score = 0
                for first_part in first_parts:
                    # Direct similarity
                    sim = self._calculate_string_similarity(first_part, patient_first)
                    max_first_score = max(max_first_score, sim)

                    # Also check against individual words in patient name
                    for patient_word in patient_first.split():
                        sim = self._calculate_string_similarity(first_part, patient_word)
                        max_first_score = max(max_first_score, sim)

                first_name_score = max_first_score

            # Score last name matching
            if last_parts and patient_last:
                max_last_score = 0
                for last_part in last_parts:
                    # Direct similarity
                    sim = self._calculate_string_similarity(last_part, patient_last)
                    max_last_score = max(max_last_score, sim)

                    # Also check against individual words in patient name
                    for patient_word in patient_last.split():
                        sim = self._calculate_string_similarity(last_part, patient_word)
                        max_last_score = max(max_last_score, sim)

                last_name_score = max_last_score

            # Calculate overall match score
            if first_name_score >= 0.6 and last_name_score >= 0.6:
                overall_score = (first_name_score + last_name_score) / 2
                matches.append((patient, overall_score, first_name_score, last_name_score))

        if matches:
            # Sort by best overall match score
            matches.sort(key=lambda x: x[1], reverse=True)
            best_match = matches[0]

            # Log the fuzzy match with detailed scores — PRN only, no names
            # (application logs may be retained/visible outside the practice's
            # own review, e.g. hosting provider log aggregation)
            self.logger.info(f"Fuzzy match found (PRN: {best_match[0].prn}) - Overall: {best_match[1]:.1%}, "
                           f"First: {best_match[2]:.1%}, Last: {best_match[3]:.1%}")

            # Return ambiguous if multiple high-scoring matches
            is_ambiguous = len([m for m in matches if m[1] >= 0.85]) > 1

            return best_match[0], is_ambiguous, best_match[1]

        # No match found
        self.logger.warning("No patient match found for a roster entry")
        return None, False, 0.0

    def _generate_invoice_lines(self, patient_df: pd.DataFrame) -> Tuple[List[InvoiceLine], float, pd.DataFrame]:
        """Generate invoice lines for a patient"""
        lines = []

        # Handle previous balance (take from first row)
        first_row = patient_df.iloc[0]
        previous_balance = float(first_row.get('previous_balance', 0))

        # Positive: owed from prior period; Negative: overpaid (credit)
        if previous_balance > 0:
            lines.append(InvoiceLine(
                service_date="",
                description="Previous Balance",
                amount=previous_balance,
                is_previous_balance=True,
                is_credit=False
            ))
        elif previous_balance < 0:
            lines.append(InvoiceLine(
                service_date="",
                description="Previous Balance (Overpaid)",
                amount=abs(previous_balance),
                is_previous_balance=True,
                is_credit=True
            ))

        # Process service lines
        for _, row in patient_df.iterrows():
            amount_due = self._calculate_amount_due(row)

            service_type = str(row.get('type_of_service', '')).strip()
            if not service_type:
                service_type = "Psychotherapy and/or Med Management"

            lines.append(InvoiceLine(
                service_date=str(row['visit_date']),
                description=service_type,
                amount=amount_due,
                is_previous_balance=False
            ))

        # Calculate total_due: negative previous_balance acts as a credit, reducing the total.
        # Deliberately NOT floored at 0 here — callers that need the true
        # (possibly negative) net balance, e.g. to detect a credit and
        # report it, would otherwise lose that information. Callers that
        # need a non-negative "amount due" for display (generate_invoices())
        # floor it themselves: a credit-balance patient still gets an
        # invoice — showing the credit as a line item and $0.00 due — not
        # a negative number.
        subtotal_copay = float(patient_df['copay'].sum()) + previous_balance
        subtotal_paid = float(patient_df['paid'].sum())
        total_due = subtotal_copay - subtotal_paid

        return lines, total_due, patient_df

    def resolve_superbill_service_lines(self, patient_df: pd.DataFrame) -> List[SuperbillServiceLine]:
        """Build superbill service lines from a patient's invoice rows,
        resolving each row's CPT code in priority order: (1) an explicit
        cpt_code column in the workbook, (2) a 5-digit code embedded in
        type_of_service (extract_embedded_cpt_code — same detection as
        _has_cpt_codes, not duplicated), (3) clinic_config's
        default_cpt_by_service_type mapping (matched on a lowercased
        type_of_service), else blank. Always meant to be reviewed/edited
        in the UI before generating — never silently trusted."""
        default_cpt_map = {k.lower(): v for k, v in self.clinic.get('default_cpt_by_service_type', {}).items()}
        lines = []
        for _, row in patient_df.iterrows():
            service_type = str(row.get('type_of_service', '')).strip()
            workbook_cpt = str(row.get('cpt_code', '')).strip()
            cpt = workbook_cpt or extract_embedded_cpt_code(service_type) or default_cpt_map.get(service_type.lower(), '')
            lines.append(SuperbillServiceLine(
                service_date=format_date_for_display(row.get('visit_date')),
                cpt_code=cpt,
                description=service_type,
                charge=float(row.get('total_amount', 0) or 0),
                payment=float(row.get('paid', 0) or 0),
            ))
        return lines

    def resolve_default_icd10_codes(self, patient_df: pd.DataFrame) -> List[str]:
        """Starting ICD-10 codes for a superbill: unique non-blank values
        from the workbook's icd10_code column if present, else
        clinic_config's default_icd10_codes. Always meant to be reviewed/
        edited in the UI, not generated without staff confirmation —
        diagnosis codes affect reimbursement and shouldn't be silently
        auto-assigned."""
        workbook_codes = sorted({
            str(c).strip() for c in patient_df.get('icd10_code', [])
            if str(c).strip()
        })
        if workbook_codes:
            return workbook_codes
        return list(self.clinic.get('default_icd10_codes', []))

    def _generate_pdf_invoice(self, patient: PatientData, lines: List[InvoiceLine],
                              total_due: float, patient_df: pd.DataFrame, output_path: Path,
                              notice_level: int = NOTICE_LEVEL_NORMAL):
        """Generate PDF invoice, automatically fitting content to a single page."""
        try:
            has_cpt = self._has_cpt_codes(patient_df)
            previous_balance = float(patient_df.iloc[0].get('previous_balance', 0))

            # Fixed margins per spec
            top_margin = 0.4 * inch
            bottom_margin = 0.6 * inch
            left_margin = 0.65 * inch
            right_margin = 0.65 * inch

            def build_pdf_bytes(p):
                """Build a complete PDF to a BytesIO buffer using the given layout params."""
                buf = BytesIO()
                doc = SimpleDocTemplate(
                    buf, pagesize=letter,
                    topMargin=top_margin, bottomMargin=bottom_margin,
                    leftMargin=left_margin, rightMargin=right_margin
                )

                story = []
                styles = getSampleStyleSheet()

                header_style = ParagraphStyle(
                    'InvHeader', parent=styles['Normal'],
                    fontSize=p['font_header'], alignment=TA_CENTER,
                    spaceAfter=3, fontName='Helvetica-Bold'
                )
                header_2_style = ParagraphStyle(
                    'InvHeader2', parent=styles['Normal'],
                    fontSize=p['font_header2'], alignment=TA_CENTER,
                    spaceAfter=2, fontName='Helvetica-Bold'
                )
                title_style = ParagraphStyle(
                    'InvTitle', parent=styles['Normal'],
                    fontSize=p['font_title'], alignment=TA_CENTER,
                    spaceAfter=p['spacer_sections'], fontName='Helvetica-Bold'
                )

                # --- Clinic header (font reduced to 10 from previous 12) ---
                story.append(Paragraph(self.clinic['clinic_name'], header_style))
                story.append(Paragraph(self.clinic['doctor_name'], header_style))
                story.append(Paragraph(self.clinic['specialty'], header_style))
                # Conditionally add EIN/NPI when CPT codes are detected
                if has_cpt:
                    story.append(Paragraph(f"EIN: {self.clinic['ein']}    NPI: {self.clinic['npi']}", header_style))
                story.append(Spacer(1, p['spacer_header']))

                # --- Contact info block (added WEBSITE line) ---
                story.append(Paragraph(f"OFFICE ADDRESS: {self.clinic['office_address']}", header_2_style))
                story.append(Paragraph(f"MAILING ADDRESS: {self.clinic['mailing_address']}", header_2_style))
                story.append(Paragraph(f"EMAIL: {self.clinic['email']}", header_2_style))
                story.append(Paragraph(f"WEBSITE: {self.clinic['website']}", header_2_style))
                story.append(Spacer(1, p['spacer_sections']))

                # --- Patient info + payment instructions side by side ---
                display_postal = patient.postal_code
                if display_postal and '.' in display_postal:
                    display_postal = display_postal.split('.')[0]

                patient_info_text = f"{patient.first_name.upper()} {patient.last_name.upper()}\n"
                if patient.address_line1:
                    patient_info_text += f"{patient.address_line1}\n"
                if patient.address_line2:
                    patient_info_text += f"{patient.address_line2}\n"
                patient_info_text += f"{patient.city}, {patient.state} {display_postal}"

                payment_info_text = (
                    "Please note we do not accept credit cards.\n"
                    f"1. Zelle {self.clinic['zelle_email']} (IRA Billing and Mgmt)\n"
                    f"2. Check payable to: {self.clinic['check_payable_to']}\n"
                    f"   {self.clinic['mailing_address']}"
                )

                combined_table = Table([[patient_info_text, payment_info_text]], colWidths=[3 * inch, 3 * inch])
                combined_table.setStyle(TableStyle([
                    ('FONTNAME', (0, 0), (0, 0), 'Helvetica-Bold'),
                    ('FONTNAME', (1, 0), (1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), p['font_body']),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 0),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                ]))
                story.append(combined_table)
                story.append(Spacer(1, p['spacer_small']))

                # --- Patient Statement title ---
                story.append(Paragraph(NOTICE_LEVEL_TITLES[notice_level], title_style))

                # --- Statement dates ---
                statement_info = [
                    ["STATEMENT DATE:", "Payment due date:"],
                    [self.statement_date.strftime('%m/%d/%Y'), self.payment_due_date.strftime('%m/%d/%Y')]
                ]
                statement_table = Table(statement_info, colWidths=[2 * inch, 2 * inch])
                statement_table.setStyle(TableStyle([
                    ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), p['font_body']),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))
                story.append(statement_table)
                story.append(Spacer(1, p['spacer_small']))

                # --- Service details table ---
                # Amount Paid column = what patient paid per visit
                # Copay/Deductible column = amount still owed per visit (copay - paid, floor 0)
                table_data = [['Service Date(s)', 'Description', 'Amount Paid', 'Copay/Deductible']]
                total_paid_display = 0.0
                total_copay_display = 0.0

                # Previous balance row(s)
                if previous_balance > 0:
                    table_data.append(['', 'Previous Balance', '$ -', f'$ {previous_balance:.2f}'])
                    total_copay_display += previous_balance
                elif previous_balance < 0:
                    # Overpayment credit goes in Amount Paid column
                    table_data.append(['', 'Previous Balance (Overpaid)', f'$ {abs(previous_balance):.2f}', '$ -'])
                    total_paid_display += abs(previous_balance)

                # Service rows
                for _, row in patient_df.iterrows():
                    paid_amount = float(row.get('paid', 0))
                    copay_amount = float(row.get('copay', 0))

                    if paid_amount > 0 or copay_amount > 0:
                        display_date = self._format_date_for_display(row['visit_date'])
                        service_type = str(row.get('type_of_service', '')).strip()
                        if not service_type:
                            service_type = "Psychotherapy and/or Med Management"

                        # Show raw copay per row; net amount due is reflected in the TOTAL row
                        table_data.append([
                            display_date,
                            service_type,
                            f'$ {paid_amount:.2f}' if paid_amount > 0 else '$ -',
                            f'$ {copay_amount:.2f}' if copay_amount > 0 else '$ -'
                        ])
                        total_paid_display += paid_amount
                        total_copay_display += copay_amount

                table_data.append(['', 'SUBTOTAL', f'$ {total_paid_display:.2f}', f'$ {total_copay_display:.2f}'])
                table_data.append(['', 'TOTAL', '', f'$ {total_due:.2f}'])

                self.logger.info(f"Table has {len(table_data)} rows")

                service_table = Table(
                    table_data,
                    colWidths=[1.5 * inch, 2.5 * inch, 1.2 * inch, 1.3 * inch]
                )
                service_table.setStyle(TableStyle([
                    # Header row
                    ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), p['font_body']),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('ALIGN', (2, 0), (3, -1), 'RIGHT'),
                    # Data rows
                    ('FONTNAME', (0, 1), (-1, -3), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -3), p['font_table']),
                    # Subtotal and total rows
                    ('FONTNAME', (0, -2), (-1, -1), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, -2), (-1, -1), p['font_body']),
                    ('LINEABOVE', (0, -2), (-1, -2), 1, colors.black),
                    ('LINEABOVE', (0, -1), (-1, -1), 2, colors.black),
                    # Grid
                    ('GRID', (0, 0), (-1, -3), 0.5, colors.black),
                    ('BOX', (0, 0), (-1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    # Dynamic padding
                    ('TOPPADDING', (0, 0), (-1, -1), p['table_top_pad']),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), p['table_bot_pad']),
                ]))
                story.append(service_table)
                story.append(Spacer(1, p['spacer_small']))

                # --- Amount due section ---
                amount_section = [
                    ["YOUR PORTION DUE:", "AMOUNT ENCLOSED:"],
                    [f"${total_due:.2f}", ""]
                ]
                amount_table = Table(amount_section, colWidths=[2 * inch, 2 * inch])
                amount_table.setStyle(TableStyle([
                    ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), p['font_body'] + 1),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('BOX', (0, 0), (0, 1), 1, colors.black),
                    ('BOX', (1, 0), (1, 1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ]))
                story.append(amount_table)
                story.append(Spacer(1, p['spacer_sections']))

                # --- Provider signature (always on same page as invoice table) ---
                signature_style = ParagraphStyle(
                    'InvSig', parent=styles['Normal'],
                    fontSize=p['font_body'] + 1, alignment=TA_RIGHT, fontName='Helvetica-Bold'
                )
                story.append(Paragraph("_________________________________", signature_style))
                story.append(Spacer(1, 8))
                story.append(Paragraph(f"Provider Signature - {self.clinic['provider_name_for_signature']}", signature_style))

                doc.build(story, onFirstPage=self.add_optimized_footer, onLaterPages=self.add_optimized_footer)
                return buf.getvalue()

            # Try each layout tier until content fits on one page
            last_pdf_bytes = None
            for tier_idx, p in enumerate(self._LAYOUT_TIERS):
                pdf_bytes = build_pdf_bytes(p)
                last_pdf_bytes = pdf_bytes
                page_count = self._count_pdf_pages(pdf_bytes)
                self.logger.info(f"Layout tier {tier_idx}: {page_count} page(s)")
                if page_count <= 1:
                    break

            output_path.write_bytes(last_pdf_bytes)
            self.logger.info("Generated PDF invoice")  # path omitted: contains patient name

        except Exception as e:
            self.logger.error(f"Error generating PDF invoice: {e}")
            raise

    def add_optimized_footer(self, canvas, doc):
        """Add two-line footer centered at the bottom of each page, plus an
        optional QR code in the bottom-right corner (config: show_qr /
        qr_image_path / qr_content — see qr_code.py). The footer text is
        horizontally centered and leaves clear space on the right at this
        length, so a ~0.9in QR in the corner doesn't overlap it."""
        canvas.saveState()
        font_size = 8
        canvas.setFont("Helvetica", font_size)

        line1 = f"If you have questions regarding your bill, please contact us at {self.clinic['phone']}."
        line2 = f"For current pricing, please visit: {self.clinic['pricing_page_url']}"

        page_width = letter[0]

        line1_width = canvas.stringWidth(line1, "Helvetica", font_size)
        line2_width = canvas.stringWidth(line2, "Helvetica", font_size)

        x1 = (page_width - line1_width) / 2
        x2 = (page_width - line2_width) / 2

        canvas.drawString(x1, 0.55 * inch, line1)
        canvas.drawString(x2, 0.35 * inch, line2)

        qr_bytes = resolve_qr_image_bytes(self.clinic)
        if qr_bytes:
            qr_size = 0.9 * inch
            qr_x = page_width - (0.65 * inch) - qr_size
            qr_y = 0.2 * inch
            canvas.drawImage(ImageReader(BytesIO(qr_bytes)), qr_x, qr_y,
                              width=qr_size, height=qr_size, mask='auto')

        canvas.restoreState()

    def _generate_cover_letter(self, patient: PatientData, template_file: str, output_path: Path):
        """Generate cover letter DOCX from template"""
        try:
            # Check if file already exists and is potentially locked
            if output_path.exists():
                try:
                    output_path.unlink()
                except OSError as e:
                    self.logger.warning(f"Could not delete existing cover letter file: {e}")
                    # Try alternative filename
                    counter = 1
                    while True:
                        new_path = output_path.parent / f"{output_path.stem}_{counter}.docx"
                        if not new_path.exists():
                            output_path = new_path
                            break
                        counter += 1
                        if counter > 10:  # Prevent infinite loop
                            raise OSError(f"Cannot create unique filename for {output_path}")

            # Load template
            doc = Document(template_file)

            # Clean postal code for display
            display_postal = patient.postal_code
            if display_postal and '.' in display_postal:
                display_postal = display_postal.split('.')[0]

            # Replacement mappings - handle empty values gracefully.
            # Keys come from the shared REQUIRED_TEMPLATE_PLACEHOLDERS list
            # (invoice_models.py) so the generator and the template validator
            # never drift apart.
            placeholder_values = [
                patient.first_name or '',
                patient.last_name or '',
                f"{patient.first_name} {patient.last_name}".strip(),
                patient.address_line1 or '',
                patient.address_line2 or '',
                patient.city or '',
                patient.state or '',
                display_postal or '',
                patient.prn or '',
            ]
            replacements = dict(zip(REQUIRED_TEMPLATE_PLACEHOLDERS, placeholder_values))
            self._replace_placeholders_in_docx(doc, replacements)

            # Handle address formatting - clean up empty address line 2
            for paragraph in doc.paragraphs:
                text = paragraph.text
                if '[Address Line 2]' in text and not patient.address_line2:
                    text = text.replace(', [Address Line 2]', '')
                    text = text.replace('[Address Line 2], ', '')
                    text = text.replace('[Address Line 2]', '')
                    paragraph.clear()
                    paragraph.add_run(text)

            doc.save(str(output_path))
            self.logger.info("Generated cover letter")  # path omitted: contains patient name

        except Exception as e:
            self.logger.error(f"Error generating cover letter: {e}")
            raise

    @staticmethod
    def _replace_placeholders_in_docx(doc, replacements: Dict[str, str]) -> None:
        """Replace [Bracket] placeholders with values throughout a docx's
        paragraphs and table cells, preserving run formatting where a
        placeholder sits entirely within one run. Shared by
        _generate_cover_letter() and _generate_notice_letter() so the
        replacement logic isn't duplicated between them."""
        def replace_text_in_paragraph(paragraph, old_text, new_text):
            if old_text in paragraph.text:
                for run in paragraph.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)

                full_text = paragraph.text
                if old_text in full_text:
                    new_text_full = full_text.replace(old_text, new_text)
                    paragraph.clear()
                    paragraph.add_run(new_text_full)

        for paragraph in doc.paragraphs:
            for placeholder, value in replacements.items():
                replace_text_in_paragraph(paragraph, placeholder, value)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for placeholder, value in replacements.items():
                            replace_text_in_paragraph(paragraph, placeholder, value)

    def _generate_notice_letter(self, patient: PatientData, notice_level: int,
                                 amount_due: float, output_path: Path):
        """Generate the 2nd/Final Notice reminder letter DOCX — replaces the
        normal cover letter for a patient the user chose to send a notice
        to instead of skipping (see generate_invoices()'s
        notice_patient_levels param). Uses NOTICE_TEMPLATE_FILES, which
        already contain the clinic's real letter wording with [Full Name] /
        [Date] / [Amount] placeholders (converted from the original Word
        mail-merge .doc templates)."""
        try:
            if output_path.exists():
                output_path.unlink()

            template_file = self.NOTICE_TEMPLATE_FILES[notice_level]
            doc = Document(str(template_file))

            full_name = f"{patient.first_name} {patient.last_name}".strip()
            date_str = f"{self.statement_date:%B} {self.statement_date.day}, {self.statement_date.year}"
            replacements = {
                '[Full Name]': full_name,
                '[Date]': date_str,
                '[Amount]': f"${amount_due:,.2f}",
            }
            self._replace_placeholders_in_docx(doc, replacements)

            doc.save(str(output_path))
            self.logger.info("Generated notice letter")  # path omitted: contains patient name

        except Exception as e:
            self.logger.error(f"Error generating notice letter: {e}")
            raise

    def _generate_envelope_pdf(self, patient: PatientData, template_file: str, output_path: Path):
        """Generate envelope PDF sized for Com-10 envelopes (4.13" x 9.5") - single page.
        NOTE: No longer called by generate_invoices(); retained for backward compatibility."""
        try:
            # Com-10 envelope dimensions
            envelope_width = 9.5 * inch
            envelope_height = 4.13 * inch
            envelope_size = (envelope_width, envelope_height)

            doc = SimpleDocTemplate(str(output_path), pagesize=envelope_size,
                                    topMargin=0.2 * inch, bottomMargin=0.2 * inch,
                                    leftMargin=0.2 * inch, rightMargin=0.2 * inch)

            story = []
            styles = getSampleStyleSheet()

            return_address_style = ParagraphStyle(
                'ReturnAddress', parent=styles['Normal'],
                fontSize=10, alignment=TA_LEFT, spaceAfter=0, spaceBefore=0,
                fontName='Helvetica', leftIndent=0, leading=12
            )

            return_address_bold_style = ParagraphStyle(
                'ReturnAddressBold', parent=styles['Normal'],
                fontSize=11, alignment=TA_LEFT, spaceAfter=1, spaceBefore=0,
                fontName='Helvetica-Bold', leftIndent=0, leading=12
            )

            delivery_address_style = ParagraphStyle(
                'DeliveryAddress', parent=styles['Normal'],
                fontSize=10, alignment=TA_LEFT, spaceAfter=0, spaceBefore=0,
                fontName='Helvetica', leftIndent=3.5 * inch, leading=12
            )

            delivery_name_style = ParagraphStyle(
                'DeliveryName', parent=styles['Normal'],
                fontSize=11, alignment=TA_LEFT, spaceAfter=2, spaceBefore=0,
                fontName='Helvetica-Bold', leftIndent=3.5 * inch, leading=12
            )

            # Return Address (very compact, top-left)
            story.append(Paragraph(f"<b>{self.clinic['clinic_name']}</b>", return_address_bold_style))
            story.append(Paragraph(self.clinic['mailing_address'], return_address_style))

            story.append(Spacer(1, 1 * inch))

            # Delivery Address (center-right)
            display_postal = patient.postal_code
            if display_postal and '.' in display_postal:
                display_postal = display_postal.split('.')[0]

            patient_full_name = f"{patient.first_name} {patient.last_name}"
            story.append(Paragraph(f"<b>{patient_full_name}</b>", delivery_name_style))

            if patient.address_line1:
                story.append(Paragraph(patient.address_line1, delivery_address_style))
            if patient.address_line2:
                story.append(Paragraph(patient.address_line2, delivery_address_style))

            city_state_zip = f"{patient.city}, {patient.state} {display_postal}"
            story.append(Paragraph(city_state_zip, delivery_address_style))

            doc.build(story)
            self.logger.info("Generated Com-10 envelope PDF")  # path omitted: contains patient name

        except Exception as e:
            self.logger.error(f"Error generating envelope PDF: {e}")
            raise

    def _generate_csv_export(self, lines: List[InvoiceLine], output_path: Path):
        """Generate CSV line items export"""
        try:
            csv_data = []

            for line in lines:
                if line.is_previous_balance:
                    if line.is_credit:
                        # Overpayment credit goes in Amount column
                        csv_data.append({
                            'Service Date': '',
                            'Description': line.description,
                            'Copay/Deductible': '',
                            'Amount': f"{line.amount:.2f}"
                        })
                    else:
                        csv_data.append({
                            'Service Date': '',
                            'Description': line.description,
                            'Copay/Deductible': f"{line.amount:.2f}",
                            'Amount': '0'
                        })
                else:
                    csv_data.append({
                        'Service Date': line.service_date,
                        'Description': line.description,
                        'Copay/Deductible': '',
                        'Amount': f"{line.amount:.2f}"
                    })

            df = pd.DataFrame(csv_data)
            df.to_csv(output_path, index=False)
            self.logger.info("Generated CSV export")  # path omitted: contains patient name

        except Exception as e:
            self.logger.error(f"Error generating CSV export: {e}")
            raise

    def _generate_summary_report_text(self, summary: ProcessingSummary,
                                       validation_report: Optional[ValidationReport] = None) -> str:
        """Render a ProcessingSummary (plus an optional pre-flight
        ValidationReport) as plain text — the same content written to
        Processing_Summary_*.txt and shown in the UI after a run."""
        lines = []
        lines.append("=" * 80)
        lines.append("PATIENT INVOICE PROCESSING SUMMARY")
        lines.append("=" * 80)
        lines.append("")

        lines.append(f"Processing Date: {summary.processing_date}")
        lines.append(f"Statement Date: {self.statement_date.strftime('%B %d, %Y')}")
        lines.append(f"Payment Due Date: {self.payment_due_date.strftime('%B %d, %Y')}")
        lines.append(f"Amount Due Strategy: {self.amount_due_strategy}")
        lines.append("")

        total_invoiced = summary.total_amount_due + summary.total_amount_paid
        lines.append("SUMMARY STATISTICS:")
        lines.append("-" * 40)
        lines.append(f"Total Patients Processed: {summary.total_processed}")
        lines.append(f"Total Patients Skipped: {summary.total_skipped}")
        lines.append(f"Total Errors: {summary.total_errors}")
        lines.append(f"Total Invoiced (billed this period): ${total_invoiced:.2f}")
        lines.append(f"Total Outstanding (amount due): ${summary.total_amount_due:.2f}")
        lines.append(f"Total Already Paid: ${summary.total_amount_paid:.2f}")
        lines.append("")

        if summary.processed_records:
            lines.append("SUCCESSFULLY PROCESSED PATIENTS:")
            lines.append("-" * 40)
            for i, record in enumerate(summary.processed_records, 1):
                if record.service_date_start and record.service_date_end:
                    start = format_date_for_display(record.service_date_start)
                    end = format_date_for_display(record.service_date_end)
                    date_range = f"{start} to {end}" if start != end else start
                else:
                    date_range = "no service dates"
                lines.append(f"{i:3d}. {record.display_name} — Service dates: {date_range} — "
                              f"Due: ${record.amount_due:.2f}, Paid: ${record.amount_paid:.2f}")
            lines.append("")

        if summary.skipped_patients:
            lines.append("SKIPPED PATIENTS:")
            lines.append("-" * 40)
            for i, (patient, reason) in enumerate(summary.skipped_patients, 1):
                lines.append(f"{i:3d}. {patient} - {reason}")
            lines.append("")

        if summary.errors:
            lines.append("ERRORS ENCOUNTERED:")
            lines.append("-" * 40)
            for i, (patient, error) in enumerate(summary.errors, 1):
                lines.append(f"{i:3d}. {patient} - {error}")
            lines.append("")

        if validation_report is not None and validation_report.issues:
            lines.append(f"VALIDATION WARNINGS (from pre-flight scan, {validation_report.generated_at}):")
            lines.append("-" * 40)
            for i, issue in enumerate(validation_report.issues, 1):
                lines.append(f"{i:3d}. [{issue.severity.upper()}] {issue.patient_name} — "
                              f"{VALIDATION_CATEGORIES.get(issue.category, issue.category)}: {issue.detail}")
            lines.append("")

        lines.append("FILES GENERATED:")
        lines.append("-" * 40)
        lines.append("For each processed patient:")
        lines.append("  - PDF and/or Excel Invoice: LastName_Year_Invoice_mmddyyyy.pdf/.xlsx")
        lines.append("  - Cover Letter: LastName_Envelope.docx")
        lines.append("  - CSV Items: LastName_Year_InvoiceItems_mmddyyyy.csv")
        lines.append("")

        lines.append("=" * 80)
        lines.append("End of Summary Report")
        lines.append("=" * 80)
        return "\n".join(lines)

    def _generate_summary_report(self, summary: ProcessingSummary, output_dir: Path,
                                  validation_report: Optional[ValidationReport] = None):
        """Write the batch summary report to Processing_Summary_*.txt."""
        try:
            summary_path = output_dir / f"Processing_Summary_{self.statement_date.strftime('%Y%m%d')}.txt"
            summary_path.write_text(self._generate_summary_report_text(summary, validation_report))
            self.logger.info("Generated summary report")  # path omitted: filename embeds the statement date, not a patient name, but kept generic for consistency

        except Exception as e:
            self.logger.error(f"Error generating summary report: {e}")
            raise

    def _check_address_issue(self, patient: PatientData) -> Optional[str]:
        """Return a description of the problem if a patient's address looks
        missing or malformed, else None."""
        if not patient.address_line1.strip():
            return "Missing street address"
        if not patient.city.strip():
            return "Missing city"
        if not patient.state.strip():
            return "Missing state"
        postal = patient.postal_code.strip()
        if not postal:
            return "Missing postal code"
        if not re.match(r'^\d{5}(-\d{4})?$', postal):
            return f"Postal code '{postal}' doesn't look like a valid US ZIP"
        return None

    @staticmethod
    def _service_date_range(patient_df: pd.DataFrame) -> Optional[Tuple[str, str]]:
        """Earliest/latest parseable visit_date in a patient's rows, as ISO
        'YYYY-MM-DD' strings — the range used for duplicate-invoice overlap
        checks. None if no row has a parseable date."""
        parsed = pd.to_datetime(patient_df['visit_date'], errors='coerce').dropna()
        if parsed.empty:
            return None
        return parsed.min().strftime('%Y-%m-%d'), parsed.max().strftime('%Y-%m-%d')

    def validate_before_generation(self, roster_file: str, invoice_file: str,
                                    custom_mapping: Dict = None,
                                    run_history_db_path: Optional[Path] = None) -> ValidationReport:
        """Read-only pre-flight scan over the roster + invoice workbook —
        generates no files. Reuses load_patient_roster()/load_invoice_data()/
        _match_patient()/_generate_invoice_lines() so matching, parsing, and
        balance-calculation logic isn't duplicated; this method only adds
        the read-only inspection pass on top.

        Checks: patients missing from the roster or matched with low
        confidence, missing/malformed addresses, missing service dates,
        charges with no service description, negative (credit) balances
        — these patients are still invoiced, not skipped, but it's useful
        to see who has a credit before generating — and possible
        duplicates (an overlapping service-date range already recorded in
        run_history for the same patient).
        """
        issues: List[ValidationIssue] = []
        db_path = run_history_db_path or run_history.DEFAULT_DB_PATH

        patients = self.load_patient_roster(roster_file)
        invoice_df = self.load_invoice_data(invoice_file, custom_mapping)
        patient_groups = invoice_df.groupby('name')

        for patient_name, patient_df in patient_groups:
            patient, is_ambiguous, score = self._match_patient(patient_name, patients)

            if patient is None:
                issues.append(ValidationIssue(
                    category="unmatched_patient", severity="error",
                    patient_name=patient_name,
                    detail="No roster match found for this name.",
                ))
            elif score < self.LOW_CONFIDENCE_THRESHOLD:
                issues.append(ValidationIssue(
                    category="low_confidence_match", severity="warning",
                    patient_name=patient_name,
                    detail=f"Best match: {patient.first_name} {patient.last_name} "
                           f"(PRN: {patient.prn}), confidence {score:.0%}.",
                ))
            elif is_ambiguous:
                issues.append(ValidationIssue(
                    category="ambiguous_match", severity="warning",
                    patient_name=patient_name,
                    detail=f"Multiple roster entries matched with similar confidence; "
                           f"using {patient.first_name} {patient.last_name} (PRN: {patient.prn}).",
                ))

            if patient is not None:
                address_issue = self._check_address_issue(patient)
                if address_issue:
                    issues.append(ValidationIssue(
                        category="malformed_address", severity="warning",
                        patient_name=patient_name, detail=address_issue,
                    ))

            for _, row in patient_df.iterrows():
                visit_date = row.get('visit_date')
                if pd.isna(visit_date) or str(visit_date).strip() == '':
                    issues.append(ValidationIssue(
                        category="missing_service_date", severity="error",
                        patient_name=patient_name,
                        detail="A service line has no visit date.",
                    ))

                total_amount = float(row.get('total_amount', 0) or 0)
                copay = float(row.get('copay', 0) or 0)
                paid = float(row.get('paid', 0) or 0)
                service_type = str(row.get('type_of_service', '')).strip()
                if (total_amount > 0 or copay > 0 or paid > 0) and not service_type:
                    display_date = format_date_for_display(visit_date) or 'unknown date'
                    issues.append(ValidationIssue(
                        category="missing_description", severity="warning",
                        patient_name=patient_name,
                        detail=f"Charge/payment on {display_date} has no service description.",
                    ))

            # Same calculation generate_invoices() uses — reused directly
            # rather than recomputed, so validation can never drift from
            # real behavior. A negative net balance doesn't skip the
            # patient; it's still invoiced (showing the credit, $0.00 due).
            _lines, total_due, _df = self._generate_invoice_lines(patient_df)
            if total_due < 0:
                issues.append(ValidationIssue(
                    category="negative_balance", severity="warning",
                    patient_name=patient_name,
                    detail=f"Net balance is a credit of ${abs(total_due):.2f}. "
                           f"This patient will still be invoiced, showing the credit "
                           f"and $0.00 due.",
                ))

            date_range = self._service_date_range(patient_df)
            if date_range is not None:
                service_start, service_end = date_range
                if patient is not None:
                    key = run_history.patient_key(patient.prn, patient.first_name, patient.last_name)
                else:
                    first_name, last_name = self._parse_patient_name(patient_name)
                    key = run_history.patient_key(None, first_name, last_name)
                overlaps = run_history.find_overlapping_runs(key, service_start, service_end, db_path=db_path)
                if overlaps:
                    most_recent = overlaps[0]
                    suggested_level = run_history.suggest_notice_level(
                        overlaps, self.statement_date.strftime('%Y-%m-%d'),
                        escalate_after_days=self.NOTICE_ESCALATION_DAYS,
                    )
                    issues.append(ValidationIssue(
                        category="duplicate_invoice", severity="warning",
                        patient_name=patient_name,
                        detail=f"Already invoiced for {most_recent.service_date_start} to "
                               f"{most_recent.service_date_end} on {most_recent.invoice_date}.",
                        suggested_notice_level=suggested_level,
                    ))

        return ValidationReport(
            issues=issues,
            total_patient_groups=len(patient_groups),
            generated_at=datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        )

    @staticmethod
    def _generate_validation_report_text(report: ValidationReport) -> str:
        """Render a ValidationReport as plain text, in the same style as
        _generate_summary_report(), for export alongside Processing_Summary_*.txt."""
        lines = []
        lines.append("=" * 80)
        lines.append("PRE-FLIGHT VALIDATION REPORT")
        lines.append("=" * 80)
        lines.append("")
        lines.append(f"Generated: {report.generated_at}")
        lines.append(f"Patient groups scanned: {report.total_patient_groups}")
        lines.append(f"Errors: {report.error_count}    Warnings: {report.warning_count}")
        lines.append("")

        if not report.issues:
            lines.append("No issues found.")
        else:
            by_category: Dict[str, List[ValidationIssue]] = {}
            for issue in report.issues:
                by_category.setdefault(issue.category, []).append(issue)

            for category, label in VALIDATION_CATEGORIES.items():
                category_issues = by_category.get(category)
                if not category_issues:
                    continue
                lines.append(f"{label.upper()} ({len(category_issues)}):")
                lines.append("-" * 40)
                for issue in category_issues:
                    lines.append(f"  [{issue.severity.upper()}] {issue.patient_name}: {issue.detail}")
                lines.append("")

        lines.append("=" * 80)
        lines.append("End of Validation Report")
        lines.append("=" * 80)
        return "\n".join(lines)

    def generate_invoices(self, roster_file: str, invoice_file: str,
                          template_file: str, output_dir: str = "output",
                          custom_mapping: Dict = None, generate_csv: bool = True,
                          envelope_format: str = "docx", export_format: str = "pdf",
                          skip_patient_names: Optional[Set[str]] = None,
                          notice_patient_levels: Optional[Dict[str, int]] = None,
                          run_history_db_path: Optional[Path] = None,
                          validation_report: Optional[ValidationReport] = None):
        """Main method to generate all invoices and cover letters.

        export_format: "pdf", "excel", or "both" — controls which invoice
        file format(s) are written per patient.
        skip_patient_names: raw invoice-sheet names (matching the 'name'
        column's groupby keys) to skip entirely — e.g. patients the user
        chose not to regenerate after a duplicate-invoice warning.
        notice_patient_levels: raw invoice-sheet names mapped to
        NOTICE_LEVEL_SECOND/FINAL — patients the user chose to send an
        escalated notice to (instead of skipping or a normal invoice)
        after a duplicate-invoice warning. Changes the PDF/Excel statement
        title and replaces the normal cover letter with the matching
        reminder letter (see _generate_notice_letter()). A name in both
        skip_patient_names and notice_patient_levels is skipped — skip
        wins, since it's the more conservative choice.
        validation_report: the pre-flight ValidationReport for this same
        roster/invoice pair, if the caller already ran one — included in
        the batch summary (file + UI) for a single combined view. Purely
        informational here; doesn't affect what gets generated.
        """
        generate_pdf_invoice = export_format in ("pdf", "both")
        generate_excel_invoice_file = export_format in ("excel", "both")
        skip_patient_names = skip_patient_names or set()
        notice_patient_levels = notice_patient_levels or {}
        db_path = run_history_db_path or run_history.DEFAULT_DB_PATH
        # Initialize summary tracking
        summary = ProcessingSummary(
            processed_records=[],
            skipped_patients=[],
            errors=[],
            total_processed=0,
            total_skipped=0,
            total_errors=0,
            total_amount_due=0.0,
            total_amount_paid=0.0,
            processing_date=datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        )

        try:
            # Load data
            patients = self.load_patient_roster(roster_file)
            invoice_df = self.load_invoice_data(invoice_file, custom_mapping)

            # Create output directory
            output_path = Path(output_dir)
            output_path.mkdir(exist_ok=True)

            # Group by patient name
            patient_groups = invoice_df.groupby('name')

            for patient_name, patient_df in patient_groups:
                try:
                    if patient_name in skip_patient_names:
                        reason = "Skipped by user (duplicate invoice)"
                        self.logger.info(f"{reason}, skipping a patient")
                        summary.skipped_patients.append((patient_name, reason))
                        summary.total_skipped += 1
                        continue

                    notice_level = notice_patient_levels.get(patient_name, NOTICE_LEVEL_NORMAL)

                    # Match patient
                    patient, is_ambiguous, _match_score = self._match_patient(patient_name, patients)

                    if is_ambiguous:
                        self.logger.warning("Ambiguous roster match for a patient")

                    # Generate invoice lines. total_due can come back negative
                    # (net credit) — that patient still gets an invoice, not
                    # a skip, so they see the credit reflected in their
                    # "Previous Balance (Overpaid)" line item; only the
                    # displayed "amount due" is floored at 0, since they
                    # don't actually owe anything.
                    lines, raw_total_due, original_df = self._generate_invoice_lines(patient_df)
                    total_due = max(0, raw_total_due)

                    if raw_total_due < 0:
                        self.logger.info("Credit balance for a patient — generating invoice to show the credit")
                        summary.credit_balance_count = getattr(summary, 'credit_balance_count', 0) + 1
                    elif total_due == 0:
                        self.logger.info("Zero balance for a patient, generating invoice to show paid in full")
                        summary.zero_balance_count = getattr(summary, 'zero_balance_count', 0) + 1

                    # Create patient folder and use matched patient data for naming
                    if patient:
                        folder_name = f"{patient.last_name}_{patient.first_name}_{patient.prn}"
                        patient_display_name = f"{patient.first_name} {patient.last_name} (PRN: {patient.prn})"
                        file_patient = patient
                    else:
                        first_name, last_name = self._parse_patient_name(patient_name)
                        folder_name = f"{last_name}_{first_name}_UNKNOWN"
                        patient_display_name = f"{first_name} {last_name} (No roster match)"
                        file_patient = PatientData("", first_name, last_name, "", "", "", "", "", "")

                    folder_name = self._sanitize_filename(folder_name)
                    patient_dir = output_path / folder_name
                    patient_dir.mkdir(exist_ok=True)

                    # Generate filenames
                    year = self.statement_date.year
                    date_str = self.statement_date.strftime("%m%d%Y")
                    base_name = f"{self._sanitize_filename(file_patient.last_name)}_{year}"

                    # PDF invoice (keep date suffix)
                    if generate_pdf_invoice:
                        pdf_path = patient_dir / f"{base_name}_Invoice_{date_str}.pdf"
                        self._generate_pdf_invoice(file_patient, lines, total_due, original_df, pdf_path,
                                                    notice_level=notice_level)

                    # Excel invoice — same base filename, .xlsx extension
                    if generate_excel_invoice_file:
                        has_cpt = self._has_cpt_codes(original_df)
                        xlsx_path = patient_dir / f"{base_name}_Invoice_{date_str}.xlsx"
                        generate_excel_invoice(file_patient, lines, total_due, original_df,
                                               self.statement_date, self.payment_due_date,
                                               has_cpt, xlsx_path, notice_level=notice_level)

                    # DOCX envelope only — no date suffix, overwrites previous copy.
                    # A patient receiving an escalated notice gets the matching
                    # reminder letter instead of the normal cover letter.
                    envelope_docx_path = patient_dir / f"{self._sanitize_filename(file_patient.last_name)}_Envelope.docx"
                    if notice_level == NOTICE_LEVEL_NORMAL:
                        self._generate_cover_letter(file_patient, template_file, envelope_docx_path)
                    else:
                        self._generate_notice_letter(file_patient, notice_level, total_due, envelope_docx_path)

                    if generate_csv:
                        csv_path = patient_dir / f"{base_name}_InvoiceItems_{date_str}.csv"
                        self._generate_csv_export(lines, csv_path)

                    # Update summary
                    total_payments = float(original_df['paid'].sum())
                    date_range = self._service_date_range(original_df)
                    service_start, service_end = date_range if date_range else (None, None)
                    summary.processed_records.append(ProcessedPatientRecord(
                        display_name=patient_display_name,
                        service_date_start=service_start,
                        service_date_end=service_end,
                        amount_due=total_due,
                        amount_paid=total_payments,
                    ))
                    summary.total_processed += 1
                    summary.total_amount_due += total_due
                    summary.total_amount_paid += total_payments

                    # Record this run for future duplicate-invoice checks
                    if date_range is not None:
                        key = run_history.patient_key(file_patient.prn, file_patient.first_name, file_patient.last_name)
                        generated_filenames = [p.name for p in (
                            pdf_path if generate_pdf_invoice else None,
                            xlsx_path if generate_excel_invoice_file else None,
                            envelope_docx_path,
                            csv_path if generate_csv else None,
                        ) if p is not None]
                        run_history.record_invoice_run(
                            key, patient_display_name, service_start, service_end,
                            self.statement_date.strftime('%Y-%m-%d'), generated_filenames,
                            notice_level=notice_level, db_path=db_path,
                        )

                except Exception as e:
                    error_msg = str(e)
                    self.logger.error(f"Error processing a patient: {error_msg}")
                    summary.errors.append((patient_name, error_msg))
                    summary.total_errors += 1
                    continue

            # Generate summary report
            self._generate_summary_report(summary, output_path, validation_report=validation_report)

            self.logger.info(f"Processing complete. Generated: {summary.total_processed}, "
                             f"Skipped: {summary.total_skipped}, Errors: {summary.total_errors}")

            return summary

        except Exception as e:
            self.logger.error(f"Fatal error in invoice generation: {e}")
            return summary


def main():
    """Example usage"""
    generator = PatientInvoiceGenerator(
        amount_due_strategy="auto"
    )

    try:
        summary = generator.generate_invoices(
            roster_file="PatientListReport_active_20250912.csv",
            invoice_file="invoice_template.xlsx",
            template_file="Access Multi Letter Cover.docx",
            output_dir="output",
            generate_csv=True
        )

        if summary:
            print("Invoice generation completed successfully!")
            print(f"Processed: {summary.total_processed} patients")
            print(f"Skipped: {summary.total_skipped} patients")
            print(f"Errors: {summary.total_errors} patients")
            print(f"Total Amount Due: ${summary.total_amount_due:.2f}")
            print(f"Summary report saved in output directory")
        else:
            print("Invoice generation completed but no summary available")

    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    main()
