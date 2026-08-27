"""Regression guard for envelope generation: _generate_cover_letter(), run
against the bundled templates/Access_Multi_Letter_Cover.docx, must produce
a real #10 landscape envelope (clinic return address + patient delivery
address, no letter body) — not the old placeholder cover-letter scaffold
that once shipped in this slot. All patient data here is synthetic.
"""
import re
import zipfile
from pathlib import Path

from invoice_models import PatientData
from complete_patient_invoice_generator import PatientInvoiceGenerator
from tests.conftest import TEST_CLINIC_CONFIG

DEFAULT_TEMPLATE = Path(__file__).parent.parent / "templates" / "Access_Multi_Letter_Cover.docx"


def _all_docx_text(docx_path) -> str:
    """All literal text runs in a docx, including content inside Structured
    Document Tags (content controls) — python-docx's Document.paragraphs
    only walks w:p elements that are direct children of the body, so text
    nested inside w:sdt/w:sdtContent (as this template's return-address
    company-name line is) would silently not show up there. Reading the
    raw XML's w:t runs directly avoids that trap."""
    with zipfile.ZipFile(docx_path) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    return "".join(re.findall(r"<w:t[^>]*>(.*?)</w:t>", xml, re.DOTALL))


def _patient():
    return PatientData(prn="1001", first_name="Ravil", last_name="Asadullin", dob="",
                        address_line1="123 Elm St", address_line2="", city="Burlingame",
                        state="CA", postal_code="94010")


def _generate(tmp_path):
    gen = PatientInvoiceGenerator(amount_due_strategy="auto", statement_date="2026-07-17",
                                   clinic_config=TEST_CLINIC_CONFIG)
    out = tmp_path / "Envelope.docx"
    gen._generate_cover_letter(_patient(), str(DEFAULT_TEMPLATE), out)
    return out


class TestEnvelopePageSetup:
    """Matches the #10 envelope spec exactly: 13680x5947 DXA landscape,
    w:code="20", and specific margins — not just "landscape-ish"."""

    def test_page_size_and_orientation(self, tmp_path):
        from docx import Document
        out = _generate(tmp_path)
        section = Document(out).sections[0]
        assert section.orientation == 1  # WD_ORIENT.LANDSCAPE
        assert round(section.page_width.twips) == 13680
        assert round(section.page_height.twips) == 5947

    def test_envelope_code_present(self, tmp_path):
        out = _generate(tmp_path)
        with zipfile.ZipFile(out) as z:
            xml = z.read("word/document.xml").decode()
        match = re.search(r"<w:pgSz[^/]*/>", xml)
        assert match is not None
        assert 'w:code="20"' in match.group(0)

    def test_margins_match_spec(self, tmp_path):
        from docx import Document
        out = _generate(tmp_path)
        section = Document(out).sections[0]
        assert round(section.top_margin.twips) == 612
        assert round(section.right_margin.twips) == 173
        assert round(section.bottom_margin.twips) == 173
        assert round(section.left_margin.twips) == 720
        assert round(section.header_distance.twips) == 360
        assert round(section.footer_distance.twips) == 360


class TestEnvelopeContent:
    def test_contains_clinic_return_address(self, tmp_path):
        out = _generate(tmp_path)
        text = _all_docx_text(out)
        assert "Access Multi-Specialty Medical Clinic, Inc." in text
        assert "PO Box 351" in text
        assert "Burlingame, CA 94011" in text

    def test_return_address_company_name_not_duplicated(self, tmp_path):
        """The company-name line lives inside a Word content control
        (Structured Document Tag) in the template — a naive fix that
        inserts a second copy as a plain paragraph (because
        Document.paragraphs doesn't show the SDT-wrapped one) would
        duplicate it. Guard against that regression specifically."""
        out = _generate(tmp_path)
        text = _all_docx_text(out)
        assert text.count("Access Multi-Specialty Medical Clinic, Inc.") == 1

    def test_contains_patient_delivery_address_from_record(self, tmp_path):
        out = _generate(tmp_path)
        text = _all_docx_text(out)
        assert "Ravil Asadullin" in text
        assert "123 Elm St" in text
        assert "Burlingame, CA 94010" in text

    def test_no_letter_body_or_placeholder_scaffold(self, tmp_path):
        """An envelope carries only the two addresses — no salutation, no
        body, no leftover placeholder scaffolding."""
        out = _generate(tmp_path)
        text = _all_docx_text(out)
        for forbidden in ["PLACEHOLDER TEMPLATE", "Dear ", "Sincerely",
                           "Please find enclosed", "Patient Record Number"]:
            assert forbidden not in text, f"{forbidden!r} should not appear in an envelope"


def test_no_placeholder_template_text_in_any_committed_template():
    """Guards against the placeholder cover-letter scaffold ever shipping
    in a template again, in any of the bundled .docx templates."""
    templates_dir = Path(__file__).parent.parent / "templates"
    for docx_path in templates_dir.glob("*.docx"):
        text = _all_docx_text(docx_path)
        assert "PLACEHOLDER TEMPLATE" not in text, f"{docx_path.name} contains placeholder scaffold text"
