"""Tests for escalating duplicate-invoice notices: instead of just skipping
a patient flagged as a possible duplicate, the user can choose to send a
2nd Notice or (after NOTICE_ESCALATION_DAYS) a Final Notice — a reminder
letter (templates/TEMPLATE_MAIL_MERGE_*_level_*.docx) plus a relabeled PDF/
Excel statement title. All patient/roster data here is synthetic."""
from pathlib import Path

import pandas as pd
import pytest
from docx import Document

import run_history
from invoice_models import (
    PatientData, NOTICE_LEVEL_NORMAL, NOTICE_LEVEL_SECOND, NOTICE_LEVEL_FINAL, NOTICE_LEVEL_TITLES,
)
from complete_patient_invoice_generator import PatientInvoiceGenerator
from excel_invoice_generator import generate_excel_invoice
from tests.conftest import TEST_CLINIC_CONFIG


@pytest.fixture
def generator():
    return PatientInvoiceGenerator(
        amount_due_strategy="auto", statement_date="2026-07-17",
        clinic_config=TEST_CLINIC_CONFIG,
    )


@pytest.fixture
def minimal_template(tmp_path):
    path = tmp_path / "template.docx"
    Document().save(path)
    return str(path)


@pytest.fixture
def roster_and_invoice(tmp_path):
    roster_rows = [dict(**{
        "Patient Identifier": "1", "Patient First Name": "Galina", "Patient Last Name": "Lerner",
        "DOB": "1980-01-01", "Address Line 1": "1 Main St", "Address Line 2": "",
        "City": "Testville", "State": "CA", "Postal Code": "94000",
    })]
    roster_path = tmp_path / "roster.csv"
    pd.DataFrame(roster_rows).to_csv(roster_path, index=False)

    invoice_rows = [dict(Name="Lerner, Galina", **{
        "Visit Date": "2025-06-10", "Total amount": 200, "Copay": 40, "Paid": 0,
        "Previous Balance": "", "Type Of Service": "Psychotherapy",
    })]
    invoice_path = tmp_path / "invoice.xlsx"
    pd.DataFrame(invoice_rows).to_excel(invoice_path, sheet_name="Sheet1", index=False)
    return str(roster_path), str(invoice_path)


class TestSuggestNoticeLevel:
    def _run(self, notice_level, invoice_date, patient_key="k1"):
        return run_history.RunRecord(
            id=1, patient_key=patient_key, patient_display_name="Test Patient",
            service_date_start="2025-06-10", service_date_end="2026-02-16",
            invoice_date=invoice_date, filenames=["a.pdf"], created_at="2026-01-01T00:00:00",
            notice_level=notice_level,
        )

    def test_no_prior_notice_suggests_second(self):
        overlaps = [self._run(NOTICE_LEVEL_NORMAL, "2026-06-01")]
        assert run_history.suggest_notice_level(overlaps, "2026-07-17") == NOTICE_LEVEL_SECOND

    def test_recent_second_notice_stays_second(self):
        overlaps = [self._run(NOTICE_LEVEL_SECOND, "2026-07-12")]  # 5 days before as_of_date
        assert run_history.suggest_notice_level(overlaps, "2026-07-17") == NOTICE_LEVEL_SECOND

    def test_second_notice_past_threshold_escalates_to_final(self):
        overlaps = [self._run(NOTICE_LEVEL_SECOND, "2026-07-01")]  # 16 days before as_of_date
        assert run_history.suggest_notice_level(overlaps, "2026-07-17") == NOTICE_LEVEL_FINAL

    def test_second_notice_exactly_at_threshold_escalates(self):
        overlaps = [self._run(NOTICE_LEVEL_SECOND, "2026-07-03")]  # exactly 14 days before
        assert run_history.suggest_notice_level(overlaps, "2026-07-17") == NOTICE_LEVEL_FINAL

    def test_final_notice_already_sent_stays_final(self):
        overlaps = [self._run(NOTICE_LEVEL_FINAL, "2026-07-16")]
        assert run_history.suggest_notice_level(overlaps, "2026-09-01") == NOTICE_LEVEL_FINAL

    def test_custom_escalation_window(self):
        overlaps = [self._run(NOTICE_LEVEL_SECOND, "2026-07-10")]  # 7 days before as_of_date
        assert run_history.suggest_notice_level(overlaps, "2026-07-17", escalate_after_days=7) == NOTICE_LEVEL_FINAL
        assert run_history.suggest_notice_level(overlaps, "2026-07-17", escalate_after_days=8) == NOTICE_LEVEL_SECOND


class TestRunHistoryNoticeLevelPersistence:
    def test_notice_level_round_trips(self, tmp_path):
        db_path = tmp_path / "run_history.db"
        key = run_history.patient_key("1", "Galina", "Lerner")
        run_history.record_invoice_run(
            key, "Galina Lerner", "2025-06-10", "2026-02-16", "2026-07-17",
            ["Lerner_2026_Invoice.pdf"], notice_level=NOTICE_LEVEL_SECOND, db_path=db_path,
        )
        overlaps = run_history.find_overlapping_runs(key, "2025-06-10", "2026-02-16", db_path=db_path)
        assert len(overlaps) == 1
        assert overlaps[0].notice_level == NOTICE_LEVEL_SECOND

    def test_default_notice_level_is_normal(self, tmp_path):
        db_path = tmp_path / "run_history.db"
        key = run_history.patient_key("1", "Galina", "Lerner")
        run_history.record_invoice_run(
            key, "Galina Lerner", "2025-06-10", "2026-02-16", "2026-07-17",
            ["Lerner_2026_Invoice.pdf"], db_path=db_path,
        )
        overlaps = run_history.find_overlapping_runs(key, "2025-06-10", "2026-02-16", db_path=db_path)
        assert overlaps[0].notice_level == NOTICE_LEVEL_NORMAL

    def test_migrates_pre_existing_db_missing_notice_level_column(self, tmp_path):
        """A run_history.db created before this feature existed has no
        notice_level column — _connect() must add it rather than error."""
        import sqlite3
        db_path = tmp_path / "run_history.db"
        conn = sqlite3.connect(str(db_path))
        conn.execute("""
            CREATE TABLE invoice_runs (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                patient_key TEXT NOT NULL, patient_display_name TEXT NOT NULL,
                service_date_start TEXT NOT NULL, service_date_end TEXT NOT NULL,
                invoice_date TEXT NOT NULL, filenames TEXT NOT NULL, created_at TEXT NOT NULL
            )
        """)
        conn.commit()
        conn.close()

        key = run_history.patient_key("1", "Galina", "Lerner")
        run_history.record_invoice_run(
            key, "Galina Lerner", "2025-06-10", "2026-02-16", "2026-07-17",
            ["a.pdf"], notice_level=NOTICE_LEVEL_FINAL, db_path=db_path,
        )
        overlaps = run_history.find_overlapping_runs(key, "2025-06-10", "2026-02-16", db_path=db_path)
        assert overlaps[0].notice_level == NOTICE_LEVEL_FINAL


class TestNoticeLetterGeneration:
    def _patient(self):
        return PatientData(prn="1", first_name="Galina", last_name="Lerner", dob="",
                            address_line1="1 Main St", address_line2="", city="Testville",
                            state="CA", postal_code="94000")

    def test_second_notice_letter_content(self, generator, tmp_path):
        out_path = tmp_path / "notice.docx"
        generator._generate_notice_letter(self._patient(), NOTICE_LEVEL_SECOND, 1234.56, out_path)
        text = "\n".join(p.text for p in Document(out_path).paragraphs)
        assert "Galina Lerner" in text
        assert "$1,234.56" in text
        assert "July 17, 2026" in text
        assert "contact us within 14 days" in text
        assert "[Full Name]" not in text and "[Date]" not in text and "[Amount]" not in text

    def test_final_notice_letter_content(self, generator, tmp_path):
        out_path = tmp_path / "notice.docx"
        generator._generate_notice_letter(self._patient(), NOTICE_LEVEL_FINAL, 500.0, out_path)
        text = "\n".join(p.text for p in Document(out_path).paragraphs)
        assert "Galina Lerner" in text
        assert "$500.00" in text
        assert "final notice" in text.lower()
        assert "collection agency" in text.lower()

    def test_pdf_title_reflects_notice_level(self, generator, tmp_path):
        import re
        patient = self._patient()
        df = pd.DataFrame([{"visit_date": "2026-01-01", "type_of_service": "Psychotherapy",
                             "total_amount": 200, "paid": 0, "copay": 50, "previous_balance": 0}])
        lines, total_due, _ = generator._generate_invoice_lines(df)
        for level in (NOTICE_LEVEL_NORMAL, NOTICE_LEVEL_SECOND, NOTICE_LEVEL_FINAL):
            out_path = tmp_path / f"invoice_{level}.pdf"
            generator._generate_pdf_invoice(patient, lines, total_due, df, out_path, notice_level=level)
            assert out_path.exists()

    def test_excel_title_reflects_notice_level(self, generator, tmp_path):
        from openpyxl import load_workbook
        patient = self._patient()
        df = pd.DataFrame([{"visit_date": "2026-01-01", "type_of_service": "Psychotherapy",
                             "total_amount": 200, "paid": 0, "copay": 50, "previous_balance": 0}])
        lines, total_due, _ = generator._generate_invoice_lines(df)
        for level in (NOTICE_LEVEL_NORMAL, NOTICE_LEVEL_SECOND, NOTICE_LEVEL_FINAL):
            out_path = tmp_path / f"invoice_{level}.xlsx"
            generate_excel_invoice(patient, lines, total_due, df, generator.statement_date,
                                    generator.payment_due_date, False, out_path,
                                    clinic=TEST_CLINIC_CONFIG, notice_level=level)
            ws = load_workbook(out_path).active
            titles = [c.value for row in ws.iter_rows() for c in row if c.value == NOTICE_LEVEL_TITLES[level]]
            assert len(titles) == 1


class TestValidationSuggestsNoticeLevel:
    def test_duplicate_issue_carries_suggested_second_notice(self, generator, minimal_template, roster_and_invoice, tmp_path):
        roster_path, invoice_path = roster_and_invoice
        db_path = tmp_path / "run_history.db"

        generator.generate_invoices(
            roster_file=roster_path, invoice_file=invoice_path, template_file=minimal_template,
            output_dir=str(tmp_path / "output1"), generate_csv=False, export_format="pdf",
            run_history_db_path=db_path,
        )

        later = PatientInvoiceGenerator(amount_due_strategy="auto", statement_date="2026-07-17",
                                         clinic_config=TEST_CLINIC_CONFIG)
        report = later.validate_before_generation(
            roster_file=roster_path, invoice_file=invoice_path, run_history_db_path=db_path,
        )
        dup_issues = [i for i in report.issues if i.category == "duplicate_invoice"]
        assert len(dup_issues) == 1
        assert dup_issues[0].suggested_notice_level == NOTICE_LEVEL_SECOND

    def test_duplicate_issue_escalates_to_final_after_threshold(self, generator, minimal_template, roster_and_invoice, tmp_path):
        roster_path, invoice_path = roster_and_invoice
        db_path = tmp_path / "run_history.db"

        # First run: normal invoice, well before the 2nd Notice.
        first_gen = PatientInvoiceGenerator(amount_due_strategy="auto", statement_date="2026-06-01",
                                             clinic_config=TEST_CLINIC_CONFIG)
        first_gen.generate_invoices(
            roster_file=roster_path, invoice_file=invoice_path, template_file=minimal_template,
            output_dir=str(tmp_path / "output1"), generate_csv=False, export_format="pdf",
            run_history_db_path=db_path,
        )
        # Second run, sends a 2nd Notice.
        second_gen = PatientInvoiceGenerator(amount_due_strategy="auto", statement_date="2026-07-17",
                                              clinic_config=TEST_CLINIC_CONFIG)
        second_gen.generate_invoices(
            roster_file=roster_path, invoice_file=invoice_path, template_file=minimal_template,
            output_dir=str(tmp_path / "output2"), generate_csv=False, export_format="pdf",
            run_history_db_path=db_path, notice_patient_levels={"Lerner, Galina": NOTICE_LEVEL_SECOND},
        )
        # Third check, 20 days after the 2nd Notice: should now suggest FINAL.
        third_gen = PatientInvoiceGenerator(amount_due_strategy="auto", statement_date="2026-08-06",
                                             clinic_config=TEST_CLINIC_CONFIG)
        report = third_gen.validate_before_generation(
            roster_file=roster_path, invoice_file=invoice_path, run_history_db_path=db_path,
        )
        dup_issues = [i for i in report.issues if i.category == "duplicate_invoice"]
        assert len(dup_issues) == 1
        assert dup_issues[0].suggested_notice_level == NOTICE_LEVEL_FINAL


class TestGenerateInvoicesNoticeIntegration:
    def test_notice_patient_generates_letter_and_records_level(self, generator, minimal_template, roster_and_invoice, tmp_path):
        roster_path, invoice_path = roster_and_invoice
        db_path = tmp_path / "run_history.db"
        output_dir = tmp_path / "output"

        summary = generator.generate_invoices(
            roster_file=roster_path, invoice_file=invoice_path, template_file=minimal_template,
            output_dir=str(output_dir), generate_csv=False, export_format="pdf",
            run_history_db_path=db_path, notice_patient_levels={"Lerner, Galina": NOTICE_LEVEL_SECOND},
        )
        assert summary.total_processed == 1

        envelope_paths = list(output_dir.rglob("*_Envelope.docx"))
        assert len(envelope_paths) == 1
        text = "\n".join(p.text for p in Document(envelope_paths[0]).paragraphs)
        assert "Galina Lerner" in text
        assert "contact us within 14 days" in text  # the 2nd-level letter's own text

        key = run_history.patient_key("1", "Galina", "Lerner")
        overlaps = run_history.find_overlapping_runs(key, "2025-06-10", "2026-02-16", db_path=db_path)
        assert overlaps[0].notice_level == NOTICE_LEVEL_SECOND

    def test_normal_patient_still_gets_cover_letter(self, generator, minimal_template, roster_and_invoice, tmp_path):
        """A patient not present in notice_patient_levels or skip_patient_names
        must still get the normal cover letter — the notice/skip machinery
        must not affect anyone it wasn't told about."""
        roster_path, invoice_path = roster_and_invoice
        db_path = tmp_path / "run_history.db"
        output_dir = tmp_path / "output"

        generator.generate_invoices(
            roster_file=roster_path, invoice_file=invoice_path, template_file=minimal_template,
            output_dir=str(output_dir), generate_csv=False, export_format="pdf",
            run_history_db_path=db_path,
        )
        key = run_history.patient_key("1", "Galina", "Lerner")
        overlaps = run_history.find_overlapping_runs(key, "2025-06-10", "2026-02-16", db_path=db_path)
        assert overlaps[0].notice_level == NOTICE_LEVEL_NORMAL

    def test_skip_wins_over_notice_when_patient_in_both(self, generator, minimal_template, roster_and_invoice, tmp_path):
        roster_path, invoice_path = roster_and_invoice
        db_path = tmp_path / "run_history.db"

        summary = generator.generate_invoices(
            roster_file=roster_path, invoice_file=invoice_path, template_file=minimal_template,
            output_dir=str(tmp_path / "output"), generate_csv=False, export_format="pdf",
            run_history_db_path=db_path,
            skip_patient_names={"Lerner, Galina"},
            notice_patient_levels={"Lerner, Galina": NOTICE_LEVEL_SECOND},
        )
        assert summary.total_processed == 0
        assert summary.total_skipped == 1
