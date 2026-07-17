#!/usr/bin/env python3
"""
Shared data structures and formatting helpers for invoice generation.
Used by both the PDF (complete_patient_invoice_generator.py) and Excel
(excel_invoice_generator.py) invoice generators so neither duplicates them.
"""
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import List, Tuple, Optional, Union, BinaryIO
import logging

import pandas as pd

# Placeholders the cover-letter DOCX template must contain, in the same
# order _generate_cover_letter() fills them in. Single source of truth for
# both the replacement logic and template validation.
REQUIRED_TEMPLATE_PLACEHOLDERS = [
    '[First Name]', '[Last Name]', '[Full Name]',
    '[Address Line 1]', '[Address Line 2]', '[City]', '[State]',
    '[Postal Code]', '[Patient Record Number]',
]


@dataclass
class PatientData:
    """Patient information from roster"""
    prn: str
    first_name: str
    last_name: str
    dob: str
    address_line1: str
    address_line2: str
    city: str
    state: str
    postal_code: str


@dataclass
class InvoiceLine:
    """Individual service line from invoice"""
    service_date: str
    description: str
    amount: float
    is_previous_balance: bool = False
    is_credit: bool = False  # True for negative previous balance (overpayment shown as credit)


@dataclass
class ProcessingSummary:
    """Summary of processing results"""
    processed_patients: List[str]
    skipped_patients: List[Tuple[str, str]]  # (name, reason)
    errors: List[Tuple[str, str]]  # (patient_name, error)
    total_processed: int
    total_skipped: int
    total_errors: int
    total_amount_due: float
    processing_date: str


# Pre-flight validation issue categories, in the order the validation panel
# groups them.
VALIDATION_CATEGORIES = {
    "unmatched_patient": "No roster match",
    "low_confidence_match": "Low-confidence roster match",
    "ambiguous_match": "Ambiguous roster match",
    "malformed_address": "Missing or malformed address",
    "missing_service_date": "Missing service date",
    "missing_description": "Charge with no service description",
    "negative_balance": "Credit / negative balance",
}


@dataclass
class ValidationIssue:
    """A single pre-flight data-quality finding for one patient group."""
    category: str  # key into VALIDATION_CATEGORIES
    severity: str  # "error" or "warning"
    patient_name: str  # raw name as it appears in the invoice workbook
    detail: str  # human-readable explanation, no PHI beyond the patient's own name


@dataclass
class ValidationReport:
    """Result of a pre-flight scan over the roster + invoice workbook,
    before any files are generated."""
    issues: List[ValidationIssue]
    total_patient_groups: int
    generated_at: str

    @property
    def error_count(self) -> int:
        return sum(1 for i in self.issues if i.severity == "error")

    @property
    def warning_count(self) -> int:
        return sum(1 for i in self.issues if i.severity == "warning")


def format_date_for_display(date_value, logger: Optional[logging.Logger] = None) -> str:
    """Format date value to MM/DD/YYYY format. Shared by PDF and Excel invoice generators."""
    if pd.isna(date_value) or date_value == '' or date_value is None:
        return ''

    try:
        # Convert to string first
        date_str = str(date_value)

        # Try to parse as datetime if it's not already
        if isinstance(date_value, str):
            # Handle different possible input formats
            for fmt in ['%Y-%m-%d', '%m/%d/%Y', '%m-%d-%Y', '%Y/%m/%d']:
                try:
                    parsed_date = datetime.strptime(date_str.split()[0], fmt)  # Take only date part
                    return parsed_date.strftime('%m/%d/%Y')
                except ValueError:
                    continue
        elif hasattr(date_value, 'strftime'):
            # If it's already a datetime object
            return date_value.strftime('%m/%d/%Y')

        # If we can't parse it, try to extract date from string
        date_part = date_str.split()[0] if ' ' in date_str else date_str

        # Try pandas to_datetime as last resort
        parsed_date = pd.to_datetime(date_part, errors='coerce')
        if not pd.isna(parsed_date):
            return parsed_date.strftime('%m/%d/%Y')

        # If all else fails, return the original string
        return date_part

    except Exception as e:
        if logger:
            logger.warning(f"Could not format date '{date_value}': {e}")
        return str(date_value) if date_value else ''


def validate_cover_letter_template(template: Union[str, Path, BinaryIO]) -> List[str]:
    """Open a cover letter .docx (path or file-like object) and report which
    of REQUIRED_TEMPLATE_PLACEHOLDERS are missing from its text. Returns an
    empty list if the template is valid. Raises if the file isn't a readable
    .docx at all.
    """
    from docx import Document  # local import: keeps python-docx optional for callers that only need models

    doc = Document(template)

    text_chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text_chunks.extend(p.text for p in cell.paragraphs)

    full_text = "\n".join(text_chunks)
    return [placeholder for placeholder in REQUIRED_TEMPLATE_PLACEHOLDERS if placeholder not in full_text]
